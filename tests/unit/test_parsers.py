from pathlib import Path

import pytest
from docx import Document

from docforge_core.parsers import (
    DocxParser,
    HtmlParser,
    ImageParser,
    MarkdownParser,
    ParsedChunk,
    PdfParser,
    TextParser,
)


def test_text_parser_parses_txt(tmp_path: Path) -> None:
    path = tmp_path / "note.txt"
    path.write_text("第一段\n\n第二段", encoding="utf-8")

    chunks = TextParser().parse(path)

    assert [chunk.text for chunk in chunks] == ["第一段", "第二段"]
    assert chunks[0].metadata["parser"] == "text"


def test_markdown_parser_parses_md(tmp_path: Path) -> None:
    path = tmp_path / "doc.md"
    path.write_text("# 标题\n正文\n\n## 小节\n内容", encoding="utf-8")

    chunks = MarkdownParser().parse(path)

    assert len(chunks) >= 2
    assert chunks[0].metadata["parser"] == "markdown"


def test_html_parser_parses_html(tmp_path: Path) -> None:
    path = tmp_path / "page.html"
    path.write_text(
        "<html><script>bad()</script><body><h1>标题</h1><p>正文内容</p></body></html>",
        encoding="utf-8",
    )

    chunks = HtmlParser().parse(path)

    assert [chunk.text for chunk in chunks] == ["标题", "正文内容"]
    assert chunks[0].metadata["parser"] == "html"


def test_docx_parser_parses_paragraphs(tmp_path: Path) -> None:
    path = tmp_path / "doc.docx"
    document = Document()
    document.add_paragraph("段落内容")
    document.save(str(path))

    chunks = DocxParser().parse(path)

    assert any(chunk.text == "段落内容" for chunk in chunks)
    assert any(chunk.metadata["block_type"] == "paragraph" for chunk in chunks)


def test_docx_parser_parses_tables(tmp_path: Path) -> None:
    path = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    document.save(str(path))

    chunks = DocxParser().parse(path)

    assert any("字段 | 值" in chunk.text for chunk in chunks)
    assert any(chunk.metadata["block_type"] == "table" for chunk in chunks)


def test_pdf_parser_missing_and_unsupported_file_behavior(tmp_path: Path) -> None:
    parser = PdfParser()

    with pytest.raises(FileNotFoundError):
        parser.parse(tmp_path / "missing.pdf")

    path = tmp_path / "bad.txt"
    path.write_text("not pdf", encoding="utf-8")
    with pytest.raises(ValueError):
        parser.parse(path)


def test_image_parser_returns_registration_chunk_without_ocr(tmp_path: Path) -> None:
    path = tmp_path / "screen.png"
    path.write_bytes(b"fake-image")

    chunks = ImageParser().parse(path)

    assert len(chunks) == 1
    assert chunks[0].text == "图片文件已登记，视觉解析将在后续 Sprint 实现。"
    assert chunks[0].metadata["parser"] == "image"
    assert chunks[0].metadata["visual_parse_status"] == "pending"


def test_parser_supports_recognizes_suffixes() -> None:
    assert TextParser().supports(Path("a.txt"))
    assert MarkdownParser().supports(Path("a.markdown"))
    assert HtmlParser().supports(Path("a.htm"))
    assert DocxParser().supports(Path("a.docx"))
    assert PdfParser().supports(Path("a.pdf"))
    assert ImageParser().supports(Path("a.webp"))
    assert not ImageParser().supports(Path("a.txt"))


def test_parsed_chunk_metadata_defaults_are_independent() -> None:
    first = ParsedChunk(text="first")
    second = ParsedChunk(text="second")

    first.metadata["key"] = "value"

    assert second.metadata == {}
    assert first.metadata is not second.metadata
