"""Read-only DOCX acceptance checks for product-level validation."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from docforge_core.domain.schemas import DocForgeState
from docforge_core.exporters.docx_exporter import INTERNAL_FIELD_TOKENS


@dataclass(frozen=True, slots=True)
class DocxAcceptanceReport:
    """Acceptance result for one DOCX file."""

    docx_path: str
    passed: bool
    issues: list[str] = field(default_factory=list)
    checked_tokens: list[str] = field(default_factory=list)


class DocxAcceptanceChecker:
    """Validate a generated DOCX without modifying it."""

    EXTRA_INTERNAL_TOKENS = (
        "export_manifest",
        "revision_trace",
        "source_draft_hash",
        "source_figure_slots_hash",
        "source_quality_gate_report_hash",
        "draft_v1",
        "draft_v2",
        "draft_v3",
        "草稿版本",
    )

    def check_normal_docx(
        self,
        docx_path: Path,
        state: DocForgeState | None = None,
        *,
        raw_quotes: list[str] | None = None,
    ) -> DocxAcceptanceReport:
        return self._check(docx_path, state, raw_quotes=raw_quotes or [], risk=False)

    def check_risk_docx(
        self,
        docx_path: Path,
        state: DocForgeState | None = None,
        *,
        raw_quotes: list[str] | None = None,
    ) -> DocxAcceptanceReport:
        return self._check(docx_path, state, raw_quotes=raw_quotes or [], risk=True)

    def _check(
        self,
        docx_path: Path,
        state: DocForgeState | None,
        *,
        raw_quotes: list[str],
        risk: bool,
    ) -> DocxAcceptanceReport:
        issues: list[str] = []
        try:
            text = self._read_text(docx_path)
        except Exception as exc:
            return DocxAcceptanceReport(
                docx_path=str(docx_path),
                passed=False,
                issues=[f"DOCX 无法打开: {exc}"],
                checked_tokens=[],
            )

        if not docx_path.exists():
            issues.append("DOCX 文件不存在")
        if risk:
            self._check_risk_content(text, issues)
        else:
            self._check_normal_content(text, state, issues)
        self._check_internal_leaks(text, raw_quotes, issues)
        return DocxAcceptanceReport(
            docx_path=str(docx_path),
            passed=not issues,
            issues=issues,
            checked_tokens=list(dict.fromkeys([*INTERNAL_FIELD_TOKENS, *self.EXTRA_INTERNAL_TOKENS])),
        )

    @staticmethod
    def _read_text(path: Path) -> str:
        document = Document(path)
        texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        texts.append(cell.text)
        return "\n".join(texts)

    @staticmethod
    def _check_normal_content(
        text: str,
        state: DocForgeState | None,
        issues: list[str],
    ) -> None:
        if "软件著作权文档" not in text:
            issues.append("DOCX 缺少软件著作权文档标题")
        product_name = ""
        version = ""
        if state is not None:
            product_name = state.target_product_name
            if state.frozen_doc_plan is not None:
                identity = state.frozen_doc_plan.software_identity
                product_name = str(
                    identity.get("target_product_name") or product_name
                )
                version = str(
                    identity.get("version") or identity.get("software_version") or ""
                )
            version = version or str(
                state.output_requirements.get("version")
                or state.output_requirements.get("software_version")
                or ""
            )
        if product_name and product_name not in text:
            issues.append("DOCX 缺少软件名称")
        if version and version not in text:
            issues.append("DOCX 缺少软件版本")
        if "核心功能" not in text and "软件概述" not in text:
            issues.append("DOCX 缺少章节标题")
        if len(text.strip()) < 80:
            issues.append("DOCX 正文内容过少")
        if state is not None and state.figure_slots_ref and "此处建议插入" not in text:
            issues.append("DOCX 缺少图片占位符")

    @staticmethod
    def _check_risk_content(text: str, issues: list[str]) -> None:
        if "风险版文档" not in text:
            issues.append("风险版 DOCX 缺少风险版标题")
        if "人工复核" not in text and "不建议直接作为正式提交版本" not in text:
            issues.append("风险版 DOCX 缺少人工复核提示")
        if "blocker_count" not in text:
            issues.append("风险版 DOCX 缺少 blocker_count 摘要")
        if "major_count" not in text:
            issues.append("风险版 DOCX 缺少 major_count 摘要")

    def _check_internal_leaks(
        self,
        text: str,
        raw_quotes: list[str],
        issues: list[str],
    ) -> None:
        lower_text = text.lower()
        if "ev_" in lower_text:
            issues.append("DOCX 泄露 evidence_id")
        for token in [*INTERNAL_FIELD_TOKENS, *self.EXTRA_INTERNAL_TOKENS]:
            if token.lower() in lower_text:
                issues.append(f"DOCX 泄露内部字段: {token}")
        for quote in raw_quotes:
            if quote and quote in text:
                issues.append("DOCX 泄露 raw quote")
        if "finding_" in lower_text:
            issues.append("DOCX 泄露 finding_id")
        if "suggested_fix" in lower_text:
            issues.append("风险版 DOCX 泄露完整 suggested_fix")
