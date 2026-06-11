import zipfile
from pathlib import Path

import pytest
from docx import Document

from docforge_core.agents.revision_loop_service import RevisionLoopService
from docforge_core.domain.enums import (
    CapabilityType,
    ImplementationStatus,
    NextAction,
    ProductFactType,
    ValidationStatus,
    WorkflowStatus,
)
from docforge_core.domain.schemas import EvidenceSupport, ProductCapability, ProductFact
from docforge_core.exporters.docx_exporter import DocxExportService, docx_has_embedded_media
from docforge_core.gates.draft_quality_gate import DraftQualityGateService
from docforge_core.io.state_store import StateStore

from .test_audit_agent import (
    _audit_ready_state,
    _draft_path,
    _load_json,
    _run_audit,
    _sha256,
    _write_json,
)
from .test_figure_slot_planner import _figure_path
from .test_revision_loop_service import RevisionThenAuditProvider, _failed_gate_ready_state


def _quality_path(store, run_id: str, version: int = 1) -> Path:
    return store.data_dir / "runs" / run_id / "drafts" / f"quality_gate_report_v{version}.json"


def _audit_path_for_version(store, run_id: str, version: int = 1) -> Path:
    return store.data_dir / "runs" / run_id / "drafts" / f"audit_report_v{version}.json"


def _draft_path_for_version(store, run_id: str, version: int = 1) -> Path:
    return store.data_dir / "runs" / run_id / "drafts" / f"draft_v{version}.json"


def _trace_path(store, run_id: str, version: int) -> Path:
    return store.data_dir / "runs" / run_id / "drafts" / f"revision_trace_v{version}.json"


def _export_dir(store, run_id: str) -> Path:
    return store.data_dir / "runs" / run_id / "exports"


def _manifest_path(store, run_id: str) -> Path:
    return _export_dir(store, run_id) / "export_manifest.json"


def _normal_docx_path(store, run_id: str) -> Path:
    return _export_dir(store, run_id) / "软件著作权文档.docx"


def _risk_docx_path(store, run_id: str) -> Path:
    return _export_dir(store, run_id) / "软件著作权文档_风险版.docx"


def _docx_text(path: Path) -> str:
    doc = Document(path)
    texts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def _docx_package_names(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        return archive.namelist()


def _prepare_v1_passed(tmp_path):
    store, state, provider = _audit_ready_state(tmp_path)
    _rewrite_draft_content_without_raw_quotes(store, state.run_id, 1)
    _run_audit(store, state, provider)
    DraftQualityGateService(store).run(state.run_id)
    return store, store.load_state(state.run_id)


def _prepare_v2_passed(tmp_path):
    store, state = _failed_gate_ready_state(tmp_path)
    service = RevisionLoopService(store, llm_provider=RevisionThenAuditProvider())
    service.revise_current_draft(state.run_id)
    _rewrite_draft_content_without_raw_quotes(store, state.run_id, 2)
    service.audit_revised_draft(state.run_id)
    DraftQualityGateService(store).run(state.run_id, draft_version=2)
    return store, store.load_state(state.run_id)


def _prepare_v3_passed(tmp_path):
    store, state = _failed_gate_ready_state(tmp_path)
    service = RevisionLoopService(store, llm_provider=RevisionThenAuditProvider())
    service.revise_current_draft(state.run_id)
    _rewrite_draft_content_without_raw_quotes(store, state.run_id, 2)
    draft_v2_path = _draft_path_for_version(store, state.run_id, 2)
    draft_v2 = _load_json(draft_v2_path)
    draft_v2["chapters"][0]["sections"][0]["citations"] = [
        {"evidence_id": "ev_missing", "quote": "missing quote"}
    ]
    _write_json(draft_v2_path, draft_v2)
    service.audit_revised_draft(state.run_id)
    DraftQualityGateService(store).run(state.run_id, draft_version=2)
    service.revise_current_draft(state.run_id)
    _rewrite_draft_content_without_raw_quotes(store, state.run_id, 3)
    service.audit_revised_draft(state.run_id)
    DraftQualityGateService(store).run(state.run_id, draft_version=3)
    return store, store.load_state(state.run_id)


def _prepare_v3_risk(tmp_path):
    store, state = _failed_gate_ready_state(tmp_path)
    service = RevisionLoopService(store, llm_provider=RevisionThenAuditProvider())
    service.revise_current_draft(state.run_id)
    _rewrite_draft_content_without_raw_quotes(store, state.run_id, 2)
    draft_v2_path = _draft_path_for_version(store, state.run_id, 2)
    draft_v2 = _load_json(draft_v2_path)
    draft_v2["chapters"][0]["sections"][0]["citations"] = [
        {"evidence_id": "ev_missing", "quote": "missing quote"}
    ]
    _write_json(draft_v2_path, draft_v2)
    service.audit_revised_draft(state.run_id)
    DraftQualityGateService(store).run(state.run_id, draft_version=2)
    service.revise_current_draft(state.run_id)
    _rewrite_draft_content_without_raw_quotes(store, state.run_id, 3)
    draft_v3_path = _draft_path_for_version(store, state.run_id, 3)
    draft_v3 = _load_json(draft_v3_path)
    draft_v3["chapters"][0]["sections"][0]["citations"] = [
        {"evidence_id": "ev_missing", "quote": "missing quote"}
    ]
    _write_json(draft_v3_path, draft_v3)
    service.audit_revised_draft(state.run_id)
    DraftQualityGateService(store).run(state.run_id, draft_version=3)
    return store, store.load_state(state.run_id)


def _rewrite_draft_content_without_raw_quotes(store, run_id: str, version: int) -> None:
    draft_path = _draft_path_for_version(store, run_id, version)
    draft = _load_json(draft_path)
    for chapter in draft["chapters"]:
        for section in chapter["sections"]:
            section["content"] = (
                f"{section['section_title']}围绕产品当前功能进行说明，"
                "描述用户可见能力、处理流程和使用场景。"
            )
    _write_json(draft_path, draft)


def _assert_no_outputs(store, run_id: str) -> None:
    assert not _normal_docx_path(store, run_id).exists()
    assert not _risk_docx_path(store, run_id).exists()
    assert not _manifest_path(store, run_id).exists()


def test_export_normal_docx_from_v1_passed_gate(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)

    result = DocxExportService(store).export_current_docx(state.run_id)

    manifest = _load_json(_manifest_path(store, state.run_id))
    assert _normal_docx_path(store, state.run_id).exists()
    assert result.export_type == "normal"
    assert manifest["export_type"] == "normal"
    assert manifest["draft_version"] == "v1"
    assert manifest["source_revision_trace_refs"] == []
    reloaded = store.load_state(state.run_id)
    assert reloaded.workflow_status == WorkflowStatus.FINAL_EXPORTED
    assert reloaded.next_action == NextAction.STOP


def test_export_normal_docx_from_v2_passed_gate(tmp_path) -> None:
    store, state = _prepare_v2_passed(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    manifest = _load_json(_manifest_path(store, state.run_id))
    assert manifest["export_type"] == "normal"
    assert manifest["draft_version"] == "v2"
    assert manifest["source_revision_trace_refs"] == ["drafts/revision_trace_v2.json"]


def test_export_normal_docx_from_v3_passed_gate(tmp_path) -> None:
    store, state = _prepare_v3_passed(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    manifest = _load_json(_manifest_path(store, state.run_id))
    assert manifest["draft_version"] == "v3"
    assert manifest["source_revision_trace_refs"] == [
        "drafts/revision_trace_v2.json",
        "drafts/revision_trace_v3.json",
    ]


def test_export_risk_docx_from_v3_failed_gate(tmp_path) -> None:
    store, state = _prepare_v3_risk(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    manifest = _load_json(_manifest_path(store, state.run_id))
    text = _docx_text(_risk_docx_path(store, state.run_id))
    assert manifest["export_type"] == "risk"
    assert manifest["draft_version"] == "v3"
    assert manifest["unresolved_blocker_count"] + manifest["unresolved_major_count"] > 0
    assert "风险版文档" in text
    assert "audit_finding" not in text
    assert "missing quote" not in text
    assert "修复" not in text


def test_export_risk_docx_rejects_non_v3(tmp_path) -> None:
    store, state = _failed_gate_ready_state(tmp_path)
    reloaded = store.load_state(state.run_id)
    reloaded.workflow_status = WorkflowStatus.RISK_VERSION_READY
    reloaded.next_action = NextAction.EXPORT_RISK_DOCX
    store.save_state(reloaded)

    with pytest.raises(ValueError, match="风险版导出只能使用 v3"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_wrong_workflow_status(tmp_path) -> None:
    store, state, _provider = _audit_ready_state(tmp_path)

    with pytest.raises(ValueError, match="合法导出状态"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_wrong_next_action(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    state.next_action = NextAction.STOP
    store.save_state(state)

    with pytest.raises(ValueError, match="合法导出状态"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_draft_changed_after_gate(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    draft = _load_json(_draft_path(store, state.run_id))
    draft["chapters"][0]["sections"][0]["content"] += "\n篡改"
    _write_json(_draft_path(store, state.run_id), draft)

    with pytest.raises(ValueError, match="source_draft_hash|不匹配"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_figure_slots_changed_after_gate(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    figures = _load_json(_figure_path(store, state.run_id))
    figures["figure_slots"][0]["recommended_caption"] = "篡改图位"
    _write_json(_figure_path(store, state.run_id), figures)

    with pytest.raises(ValueError, match="figure_slots_hash|不匹配"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_audit_report_changed_after_gate(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    audit_path = _audit_path_for_version(store, state.run_id)
    audit = _load_json(audit_path)
    audit["findings"] = []
    _write_json(audit_path, audit)

    with pytest.raises(ValueError, match="source_audit_report_hash|不匹配"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_quality_gate_report_state_ref_mismatch(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    state.draft_quality_gate_report_ref = "drafts/quality_gate_report_v2.json"
    store.save_state(state)

    with pytest.raises(ValueError, match="缺失|不匹配"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_audit_report_state_ref_mismatch(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    state.audit_report_ref = "drafts/audit_report_v2.json"
    store.save_state(state)

    with pytest.raises(ValueError, match="state.audit_report_ref|不匹配"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_revision_trace_missing_for_v2(tmp_path) -> None:
    store, state = _prepare_v2_passed(tmp_path)
    _trace_path(store, state.run_id, 2).unlink()

    with pytest.raises(ValueError, match="revision_trace_v2|缺失"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_rejects_revision_trace_hash_mismatch(tmp_path) -> None:
    store, state = _prepare_v2_passed(tmp_path)
    trace_path = _trace_path(store, state.run_id, 2)
    trace = _load_json(trace_path)
    trace["source_draft_hash"] = "bad"
    _write_json(trace_path, trace)

    with pytest.raises(ValueError, match="revision_trace_v2 source_draft_hash|不匹配"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)


def test_export_v3_validates_revision_trace_chain(tmp_path) -> None:
    store, state = _prepare_v3_passed(tmp_path)
    trace_v2 = _load_json(_trace_path(store, state.run_id, 2))
    trace_v2["source_audit_report_hash"] = "bad"
    _write_json(_trace_path(store, state.run_id, 2), trace_v2)

    with pytest.raises(ValueError, match="revision_trace_v2 source_audit_report_hash|不匹配"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)

    store, state = _prepare_v3_passed(tmp_path / "second")
    trace_v3 = _load_json(_trace_path(store, state.run_id, 3))
    trace_v3["source_quality_gate_report_hash"] = "bad"
    _write_json(_trace_path(store, state.run_id, 3), trace_v3)
    with pytest.raises(ValueError, match="revision_trace_v3 source_quality_gate_report_hash|不匹配"):
        DocxExportService(store).export_current_docx(state.run_id)


def test_export_rejects_draft_structure_mismatch_with_section_plan(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    draft = _load_json(_draft_path(store, state.run_id))
    extra = dict(draft["chapters"][0]["sections"][0])
    extra["section_id"] = "sec_orphan"
    draft["chapters"][0]["sections"].append(extra)
    _write_json(_draft_path(store, state.run_id), draft)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_draft_hash"] = _sha256(_draft_path(store, state.run_id))
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_draft_hash"] = audit["source_draft_hash"]
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="SectionPlan|section"):
        DocxExportService(store).export_current_docx(state.run_id)


def test_export_rejects_tampered_section_title_or_path(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    draft = _load_json(_draft_path(store, state.run_id))
    draft["chapters"][0]["sections"][0]["section_title"] = "被篡改标题"
    _write_json(_draft_path(store, state.run_id), draft)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_draft_hash"] = _sha256(_draft_path(store, state.run_id))
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_draft_hash"] = audit["source_draft_hash"]
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="section_title|SectionPlan"):
        DocxExportService(store).export_current_docx(state.run_id)


def test_export_rejects_missing_required_section(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    draft = _load_json(_draft_path(store, state.run_id))
    draft["chapters"][0]["sections"].pop()
    _write_json(_draft_path(store, state.run_id), draft)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_draft_hash"] = _sha256(_draft_path(store, state.run_id))
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_draft_hash"] = audit["source_draft_hash"]
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="SectionPlan|数量"):
        DocxExportService(store).export_current_docx(state.run_id)


def test_export_rejects_chapter_order_changed(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    draft = _load_json(_draft_path(store, state.run_id))
    if len(draft["chapters"]) < 2:
        draft["chapters"].append({"chapter_id": "extra", "title": "额外一级目录", "level": 1, "sections": []})
    else:
        draft["chapters"].reverse()
    _write_json(_draft_path(store, state.run_id), draft)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_draft_hash"] = _sha256(_draft_path(store, state.run_id))
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_draft_hash"] = audit["source_draft_hash"]
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="chapter|一级目录|Outline"):
        DocxExportService(store).export_current_docx(state.run_id)


def test_export_docx_does_not_expose_internal_evidence_ids(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    text = _docx_text(_normal_docx_path(store, state.run_id))
    forbidden = [
        "ev_",
        "evidence_id",
        "source_id",
        "audit_report_id",
        "quality_gate_report_id",
        "source_draft_hash",
        "source_figure_slots_hash",
        "export_manifest",
    ]
    assert all(item not in text for item in forbidden)


def test_export_docx_does_not_include_internal_draft_version(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    text = _docx_text(_normal_docx_path(store, state.run_id)).lower()
    forbidden = [
        "草稿版本",
        "draft_version",
        "draft_v1",
        "draft_v2",
        "draft_v3",
        "revision_trace",
    ]
    assert all(item not in text for item in forbidden)


def test_export_docx_does_not_expose_raw_quote(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    text = _docx_text(_normal_docx_path(store, state.run_id))
    assert "quote" not in text
    assert "当前版本明确支持数据集管理能力" not in text


def test_export_rejects_draft_body_with_internal_field(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    draft = _load_json(_draft_path(store, state.run_id))
    draft["chapters"][0]["sections"][0]["content"] = "内部字段 evidence_id: ev_product"
    _write_json(_draft_path(store, state.run_id), draft)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_draft_hash"] = _sha256(_draft_path(store, state.run_id))
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_draft_hash"] = audit["source_draft_hash"]
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="内部字段"):
        DocxExportService(store).export_current_docx(state.run_id)


@pytest.mark.parametrize(
    "unsafe_content",
    [
        "页面ev_secret已经配置",
        "页面ev_已经配置",
        "内部source_id字段",
    ],
)
def test_export_rejects_internal_field_adjacent_to_chinese_text(
    tmp_path,
    unsafe_content: str,
) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    draft = _load_json(_draft_path(store, state.run_id))
    draft["chapters"][0]["sections"][0]["content"] = unsafe_content
    _write_json(_draft_path(store, state.run_id), draft)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_draft_hash"] = _sha256(_draft_path(store, state.run_id))
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_draft_hash"] = audit["source_draft_hash"]
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="内部字段"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)
    assert store.load_state(state.run_id).workflow_status == WorkflowStatus.DRAFT_QUALITY_GATE_PASSED


def test_export_rejects_draft_body_with_product_capability_raw_quote(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    raw_quote = "当前版本明确支持安全审计能力"
    state.product_capabilities = [
        ProductCapability(
            capability_id="cap_security_audit",
            name="安全审计",
            capability_type=CapabilityType.SYSTEM_ADMINISTRATION,
            implementation_status=ImplementationStatus.CURRENT,
            validation_status=ValidationStatus.VALIDATED,
            evidence_supports=[
                EvidenceSupport(
                    evidence_id="ev_product",
                    source_id="product_source",
                    quote=raw_quote,
                )
            ],
            confidence=0.9,
        )
    ]
    store.save_state(state)
    draft = _load_json(_draft_path(store, state.run_id))
    draft["chapters"][0]["sections"][0]["content"] = raw_quote
    _write_json(_draft_path(store, state.run_id), draft)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_draft_hash"] = _sha256(_draft_path(store, state.run_id))
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_draft_hash"] = audit["source_draft_hash"]
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="raw quote"):
        DocxExportService(store).export_current_docx(state.run_id)


def test_export_rejects_draft_body_with_product_fact_raw_quote(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    raw_quote = "系统提供组织权限配置页面"
    state.product_facts = [
        ProductFact(
            fact_type=ProductFactType.FEATURE,
            content="组织权限配置",
            source_ids=["product_source"],
            implementation_status=ImplementationStatus.CURRENT,
            validation_status=ValidationStatus.VALIDATED,
            supporting_evidence_ids=["ev_product"],
            supporting_quotes=[raw_quote],
        )
    ]
    store.save_state(state)
    draft = _load_json(_draft_path(store, state.run_id))
    draft["chapters"][0]["sections"][0]["content"] = raw_quote
    _write_json(_draft_path(store, state.run_id), draft)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_draft_hash"] = _sha256(_draft_path(store, state.run_id))
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_draft_hash"] = audit["source_draft_hash"]
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="raw quote"):
        DocxExportService(store).export_current_docx(state.run_id)


def test_export_risk_docx_does_not_dump_full_audit_report(tmp_path) -> None:
    store, state = _prepare_v3_risk(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    text = _docx_text(_risk_docx_path(store, state.run_id))
    assert "missing quote" not in text
    assert "recommendation" not in text
    assert "audit_finding" not in text


def test_export_uses_figure_slots_ref_from_gate_report(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    run_dir = store.data_dir / "runs" / state.run_id
    alt_ref = "drafts/figure_slots_alt.json"
    alt_path = run_dir / alt_ref
    figures = _load_json(_figure_path(store, state.run_id))
    figures["figure_slots"][0]["recommended_caption"] = "替代图位来源"
    _write_json(alt_path, figures)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_figure_slots_ref"] = alt_ref
    audit["source_figure_slots_hash"] = _sha256(alt_path)
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_figure_slots_ref"] = alt_ref
    gate["source_figure_slots_hash"] = _sha256(alt_path)
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    DocxExportService(store).export_current_docx(state.run_id)

    assert "替代图位来源" in _docx_text(_normal_docx_path(store, state.run_id))


def test_export_rejects_internal_field_in_figure_caption(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    run_dir = store.data_dir / "runs" / state.run_id
    alt_ref = "drafts/figure_slots_alt.json"
    alt_path = run_dir / alt_ref
    figures = _load_json(_figure_path(store, state.run_id))
    figures["figure_slots"][0]["recommended_caption"] = "用户页面 ev_secret"
    _write_json(alt_path, figures)
    audit = _load_json(_audit_path_for_version(store, state.run_id))
    audit["source_figure_slots_ref"] = alt_ref
    audit["source_figure_slots_hash"] = _sha256(alt_path)
    _write_json(_audit_path_for_version(store, state.run_id), audit)
    gate = _load_json(_quality_path(store, state.run_id))
    gate["source_figure_slots_ref"] = alt_ref
    gate["source_figure_slots_hash"] = _sha256(alt_path)
    gate["source_audit_report_hash"] = _sha256(_audit_path_for_version(store, state.run_id))
    _write_json(_quality_path(store, state.run_id), gate)

    with pytest.raises(ValueError, match="内部字段"):
        DocxExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)
    assert store.load_state(state.run_id).workflow_status == WorkflowStatus.DRAFT_QUALITY_GATE_PASSED


def test_export_does_not_insert_real_images(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    assert not docx_has_embedded_media(_normal_docx_path(store, state.run_id))
    assert all(not name.startswith("word/media/") for name in _docx_package_names(_normal_docx_path(store, state.run_id)))


def test_export_does_not_block_when_figure_slot_has_no_image(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    assert "[此处建议插入：" in _docx_text(_normal_docx_path(store, state.run_id))


def test_export_rejects_existing_docx_without_force(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    _export_dir(store, state.run_id).mkdir(parents=True, exist_ok=True)
    _normal_docx_path(store, state.run_id).write_bytes(b"existing")

    with pytest.raises(ValueError, match="已存在"):
        DocxExportService(store).export_current_docx(state.run_id)

    assert not _manifest_path(store, state.run_id).exists()


def test_export_force_overwrites_existing_docx(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    _export_dir(store, state.run_id).mkdir(parents=True, exist_ok=True)
    _normal_docx_path(store, state.run_id).write_bytes(b"existing")

    result = DocxExportService(store).export_current_docx(state.run_id, force=True)

    assert result.docx_hash == _sha256(_normal_docx_path(store, state.run_id))
    assert _manifest_path(store, state.run_id).exists()


class FailingManifestExportService(DocxExportService):
    def _build_export_manifest(self, context, docx_path, manifest_path):
        raise RuntimeError("manifest write failure")


class FailingExportStateStore(StateStore):
    def save_state(self, state):
        raise RuntimeError("save_state failure")


def test_export_manifest_write_failure_cleans_docx(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)

    with pytest.raises(RuntimeError, match="manifest write failure"):
        FailingManifestExportService(store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)
    assert store.load_state(state.run_id).workflow_status == WorkflowStatus.DRAFT_QUALITY_GATE_PASSED


def test_export_state_save_failure_cleans_outputs(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    failing_store = FailingExportStateStore(data_dir=store.data_dir)

    with pytest.raises(RuntimeError, match="save_state failure"):
        DocxExportService(failing_store).export_current_docx(state.run_id)

    _assert_no_outputs(store, state.run_id)
    assert store.load_state(state.run_id).workflow_status == WorkflowStatus.DRAFT_QUALITY_GATE_PASSED


def test_force_export_failure_restores_existing_outputs(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)
    docx_path = _normal_docx_path(store, state.run_id)
    manifest_path = _manifest_path(store, state.run_id)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    docx_path.write_bytes(b"existing docx")
    manifest_path.write_text('{"existing": true}\n', encoding="utf-8")
    failing_store = FailingExportStateStore(data_dir=store.data_dir)

    with pytest.raises(RuntimeError, match="save_state failure"):
        DocxExportService(failing_store).export_current_docx(state.run_id, force=True)

    assert docx_path.read_bytes() == b"existing docx"
    assert manifest_path.read_text(encoding="utf-8") == '{"existing": true}\n'
    assert store.load_state(state.run_id).workflow_status == WorkflowStatus.DRAFT_QUALITY_GATE_PASSED


def test_export_does_not_create_pdf_or_markdown(tmp_path) -> None:
    store, state = _prepare_v1_passed(tmp_path)

    DocxExportService(store).export_current_docx(state.run_id)

    outputs = [path.suffix.lower() for path in _export_dir(store, state.run_id).iterdir()]
    assert ".pdf" not in outputs
    assert ".md" not in outputs
