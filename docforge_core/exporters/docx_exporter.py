"""Sprint 13 safe DOCX exporter with artifact lineage validation."""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal

from docx import Document

from docforge_core.agents.outline_traversal import iter_outline_sections
from docforge_core.domain.enums import (
    AuditSeverity,
    DraftQualityGateDecision,
    ExportType,
    NextAction,
    WorkflowStatus,
)
from docforge_core.domain.schemas import (
    DocForgeState,
    DraftAuditReport,
    DraftQualityGateReport,
    ExportManifest,
    ExportResult,
    FigureSlotResult,
)
from docforge_core.io.run_paths import get_exports_dir, get_run_dir, get_state_file
from docforge_core.io.state_store import StateStore

EV_ID_PATTERN = re.compile(r"(?<![A-Za-z0-9_])ev_[A-Za-z0-9_:-]+", re.IGNORECASE)
INTERNAL_FIELD_TOKENS = (
    "evidence_id",
    "source_id",
    "source_draft_hash",
    "source_figure_slots_hash",
    "source_audit_report_hash",
    "source_quality_gate_report_hash",
    "audit_report",
    "audit_report_id",
    "quality_gate_report",
    "quality_gate_report_id",
    "revision_trace",
    "export_manifest",
    "finding_id",
    "state.json",
    "evidence_map",
    "raw quote",
    "draft_version",
    "draft_v1",
    "draft_v2",
    "draft_v3",
    "草稿版本",
)
RISK_NOTICE = (
    "该文档由系统在自动修订后仍存在关键风险项的情况下导出，"
    "仅供人工复核，不建议直接作为正式提交版本。"
)


@dataclass(frozen=True)
class ExportContext:
    export_type: Literal["normal", "risk"]
    version: int
    draft_ref: str
    audit_ref: str
    gate_ref: str
    figure_ref: str
    draft_path: Path
    audit_path: Path
    gate_path: Path
    figure_path: Path
    draft: dict[str, Any]
    figures: FigureSlotResult
    audit_report: DraftAuditReport
    gate_report: DraftQualityGateReport
    revision_trace_refs: list[str]
    revision_trace_hashes: list[str]


@dataclass(frozen=True)
class DocxExportResult:
    export_type: Literal["normal", "risk"]
    draft_version: str
    docx_path: str
    manifest_path: str
    docx_hash: str
    manifest: ExportManifest


class DocxExportService:
    """Export the current approved or risk-accepted draft to user-facing DOCX."""

    def __init__(self, state_store: StateStore | None = None) -> None:
        self.state_store = state_store or StateStore()

    def export_current_docx(self, run_id: str, force: bool = False) -> DocxExportResult:
        state = self.state_store.load_state(run_id)
        context = self._resolve_export_context(state)
        self._validate_draft_structure_against_section_plan(state, context.draft)
        generated_at = datetime.now(UTC).isoformat()
        visible_texts = self._collect_user_visible_docx_texts(context, state, generated_at)
        self._validate_user_visible_docx_texts_safe(visible_texts, state, context.draft)

        exports_dir = get_exports_dir(run_id, self.state_store.data_dir)
        exports_dir.mkdir(parents=True, exist_ok=True)
        docx_name = (
            "软件著作权文档_风险版.docx"
            if context.export_type == "risk"
            else "软件著作权文档.docx"
        )
        docx_path = exports_dir / docx_name
        manifest_path = exports_dir / "export_manifest.json"
        if not force:
            if docx_path.exists():
                raise ValueError("目标 DOCX 已存在，force=False 时不允许覆盖")
            if manifest_path.exists():
                raise ValueError("export_manifest.json 已存在，force=False 时不允许覆盖")

        state_path = get_state_file(run_id, self.state_store.data_dir)
        original_state = state_path.read_bytes()
        docx_tmp = docx_path.with_suffix(docx_path.suffix + ".tmp")
        manifest_tmp = manifest_path.with_suffix(manifest_path.suffix + ".tmp")
        original_docx = self._snapshot_file(docx_path)
        original_manifest = self._snapshot_file(manifest_path)
        for tmp in (docx_tmp, manifest_tmp):
            if tmp.exists():
                tmp.unlink()

        try:
            self._render_docx(context, state, docx_tmp, generated_at)
            self._validate_user_visible_docx_texts_safe(
                self._read_docx_visible_texts(docx_tmp),
                state,
                context.draft,
            )
            docx_tmp.replace(docx_path)
            manifest = self._build_export_manifest(context, docx_path, manifest_path)
            manifest_tmp.write_text(
                json.dumps(manifest.model_dump(mode="json"), ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
            manifest_tmp.replace(manifest_path)
            self._apply_success_state(state, context, docx_path)
            self.state_store.save_state(state)
        except Exception:
            if docx_tmp.exists():
                docx_tmp.unlink()
            if manifest_tmp.exists():
                manifest_tmp.unlink()
            self._restore_output_file(docx_path, original_docx)
            self._restore_output_file(manifest_path, original_manifest)
            self._restore_state_file(state_path, original_state)
            raise

        return DocxExportResult(
            export_type=context.export_type,
            draft_version=f"v{context.version}",
            docx_path=str(docx_path),
            manifest_path=str(manifest_path),
            docx_hash=self.sha256(docx_path),
            manifest=manifest,
        )

    def _resolve_export_context(self, state: DocForgeState) -> ExportContext:
        export_type = self._validate_export_state(state)
        version = self._draft_version_from_state_gate_ref(state)
        run_dir = get_run_dir(state.run_id, self.state_store.data_dir)
        draft_ref = f"drafts/draft_v{version}.json"
        audit_ref = f"drafts/audit_report_v{version}.json"
        gate_ref = f"drafts/quality_gate_report_v{version}.json"
        draft_path = run_dir / draft_ref
        audit_path = run_dir / audit_ref
        gate_path = run_dir / gate_ref

        draft = self._load_object(draft_path, f"draft_v{version}.json")
        audit_report = DraftAuditReport.model_validate(
            self._load_object(audit_path, f"audit_report_v{version}.json")
        )
        gate_report = DraftQualityGateReport.model_validate(
            self._load_object(gate_path, f"quality_gate_report_v{version}.json")
        )
        figure_ref = gate_report.source_figure_slots_ref or audit_report.source_figure_slots_ref
        figure_path = self._safe_run_relative_path(run_dir, figure_ref)
        figures = FigureSlotResult.model_validate(
            self._load_object(figure_path, Path(figure_ref).name)
        )

        self._validate_gate_decision_for_export(export_type, version, audit_report, gate_report)
        self._validate_lineage(
            state=state,
            version=version,
            draft_path=draft_path,
            audit_path=audit_path,
            gate_path=gate_path,
            figure_path=figure_path,
            audit_report=audit_report,
            gate_report=gate_report,
        )
        trace_refs, trace_hashes = self._validate_revision_chain(run_dir, version)

        return ExportContext(
            export_type=export_type,
            version=version,
            draft_ref=draft_ref,
            audit_ref=audit_ref,
            gate_ref=gate_ref,
            figure_ref=figure_ref,
            draft_path=draft_path,
            audit_path=audit_path,
            gate_path=gate_path,
            figure_path=figure_path,
            draft=draft,
            figures=figures,
            audit_report=audit_report,
            gate_report=gate_report,
            revision_trace_refs=trace_refs,
            revision_trace_hashes=trace_hashes,
        )

    @staticmethod
    def _validate_export_state(state: DocForgeState) -> Literal["normal", "risk"]:
        if (
            state.workflow_status == WorkflowStatus.DRAFT_QUALITY_GATE_PASSED
            and state.next_action == NextAction.EXPORT_DOCX
        ):
            return "normal"
        if (
            state.workflow_status == WorkflowStatus.RISK_VERSION_READY
            and state.next_action == NextAction.EXPORT_RISK_DOCX
        ):
            return "risk"
        raise ValueError("DocxExportService 只能从合法导出状态启动")

    @staticmethod
    def _draft_version_from_state_gate_ref(state: DocForgeState) -> int:
        if not state.draft_quality_gate_report_ref:
            raise ValueError("state.draft_quality_gate_report_ref 缺失")
        name = Path(state.draft_quality_gate_report_ref).name
        match = re.fullmatch(r"quality_gate_report_v([123])\.json", name)
        if not match:
            raise ValueError("state.draft_quality_gate_report_ref 版本不合法")
        return int(match.group(1))

    @staticmethod
    def _validate_gate_decision_for_export(
        export_type: Literal["normal", "risk"],
        version: int,
        audit_report: DraftAuditReport,
        gate_report: DraftQualityGateReport,
    ) -> None:
        if export_type == "normal":
            if not gate_report.passed:
                raise ValueError("普通版导出要求 passed quality_gate_report")
            if gate_report.next_action != NextAction.EXPORT_DOCX:
                raise ValueError("普通版导出要求 quality_gate_report next_action=EXPORT_DOCX")
            if gate_report.next_workflow_status != WorkflowStatus.DRAFT_QUALITY_GATE_PASSED:
                raise ValueError("普通版导出要求 quality_gate_report 指向 DRAFT_QUALITY_GATE_PASSED")
            if gate_report.decision != DraftQualityGateDecision.EXPORT_DOCX:
                raise ValueError("普通版导出要求 decision=export_docx")
            return

        if version != 3:
            raise ValueError("风险版导出只能使用 v3")
        if gate_report.passed:
            raise ValueError("风险版导出要求 failed quality_gate_report")
        if gate_report.decision != DraftQualityGateDecision.RISK_EXPORT_REQUIRED:
            raise ValueError("风险版导出要求 decision=risk_export_required")
        if gate_report.next_action != NextAction.EXPORT_RISK_DOCX:
            raise ValueError("风险版导出要求 next_action=EXPORT_RISK_DOCX")
        if gate_report.next_workflow_status != WorkflowStatus.RISK_VERSION_READY:
            raise ValueError("风险版导出要求 next_workflow_status=RISK_VERSION_READY")
        hard_count = sum(
            item.severity in {AuditSeverity.BLOCKER, AuditSeverity.MAJOR}
            for item in audit_report.findings
        )
        if hard_count == 0:
            raise ValueError("风险版导出要求 audit_report_v3 仍存在 blocker 或 major")

    def _validate_lineage(
        self,
        *,
        state: DocForgeState,
        version: int,
        draft_path: Path,
        audit_path: Path,
        gate_path: Path,
        figure_path: Path,
        audit_report: DraftAuditReport,
        gate_report: DraftQualityGateReport,
    ) -> None:
        expected_version = f"v{version}"
        expected_draft_ref = f"drafts/draft_v{version}.json"
        expected_audit_ref = f"drafts/audit_report_v{version}.json"
        expected_gate_ref = f"drafts/quality_gate_report_v{version}.json"

        if audit_report.draft_version != expected_version:
            raise ValueError("audit_report draft_version 不匹配")
        if audit_report.source_draft_ref != expected_draft_ref:
            raise ValueError("audit_report source_draft_ref 不匹配")
        if audit_report.source_draft_hash != self.sha256(draft_path):
            raise ValueError("audit_report source_draft_hash 不匹配")
        if audit_report.source_figure_slots_ref != gate_report.source_figure_slots_ref:
            raise ValueError("audit_report source_figure_slots_ref 不匹配")
        if audit_report.source_figure_slots_hash != self.sha256(figure_path):
            raise ValueError("audit_report source_figure_slots_hash 不匹配")
        if not audit_report.report_id.strip():
            raise ValueError("audit_report report_id 不得为空")
        if state.audit_report_ref != expected_audit_ref:
            raise ValueError("state.audit_report_ref 不匹配")
        if state.audit_report_result_id != audit_report.report_id:
            raise ValueError("state.audit_report_result_id 不匹配")

        if gate_report.draft_version != expected_version:
            raise ValueError("quality_gate_report draft_version 不匹配")
        if gate_report.source_audit_report_path != expected_audit_ref:
            raise ValueError("quality_gate_report source_audit_report_path 不匹配")
        if gate_report.source_audit_report_hash != self.sha256(audit_path):
            raise ValueError("quality_gate_report source_audit_report_hash 不匹配")
        if gate_report.audit_overall_passed != audit_report.overall_passed:
            raise ValueError("quality_gate_report audit_overall_passed 不匹配")
        if gate_report.source_draft_ref != audit_report.source_draft_ref:
            raise ValueError("quality_gate_report source_draft_ref 不匹配")
        if gate_report.source_draft_hash != audit_report.source_draft_hash:
            raise ValueError("quality_gate_report source_draft_hash 不匹配")
        if gate_report.source_figure_slots_ref != audit_report.source_figure_slots_ref:
            raise ValueError("quality_gate_report source_figure_slots_ref 不匹配")
        if gate_report.source_figure_slots_hash != audit_report.source_figure_slots_hash:
            raise ValueError("quality_gate_report source_figure_slots_hash 不匹配")
        if gate_report.source_figure_slots_hash != self.sha256(figure_path):
            raise ValueError("quality_gate_report source_figure_slots_hash 与文件不匹配")
        if state.draft_quality_gate_report_ref != expected_gate_ref:
            raise ValueError("state.draft_quality_gate_report_ref 不匹配")
        self._validate_finding_binding(audit_report, gate_report)

    @staticmethod
    def _validate_finding_binding(
        audit_report: DraftAuditReport,
        gate_report: DraftQualityGateReport,
    ) -> None:
        ids = {
            "blocker": {
                item.finding_id
                for item in audit_report.findings
                if item.severity == AuditSeverity.BLOCKER
            },
            "major": {
                item.finding_id
                for item in audit_report.findings
                if item.severity == AuditSeverity.MAJOR
            },
            "minor": {
                item.finding_id
                for item in audit_report.findings
                if item.severity == AuditSeverity.MINOR
            },
            "suggestion": {
                item.finding_id
                for item in audit_report.findings
                if item.severity == AuditSeverity.SUGGESTION
            },
        }
        if set(gate_report.blocking_finding_ids) != ids["blocker"]:
            raise ValueError("quality_gate_report blocking_finding_ids 不匹配")
        if set(gate_report.major_finding_ids) != ids["major"]:
            raise ValueError("quality_gate_report major_finding_ids 不匹配")
        if set(gate_report.minor_finding_ids) != ids["minor"]:
            raise ValueError("quality_gate_report minor_finding_ids 不匹配")
        if set(gate_report.suggestion_finding_ids) != ids["suggestion"]:
            raise ValueError("quality_gate_report suggestion_finding_ids 不匹配")
        counts = gate_report.severity_counts
        if (
            counts.blocker != len(ids["blocker"])
            or counts.major != len(ids["major"])
            or counts.minor != len(ids["minor"])
            or counts.suggestion != len(ids["suggestion"])
        ):
            raise ValueError("quality_gate_report severity_counts 不匹配")

    def _validate_revision_chain(self, run_dir: Path, version: int) -> tuple[list[str], list[str]]:
        trace_refs: list[str] = []
        trace_hashes: list[str] = []
        if version == 1:
            return trace_refs, trace_hashes
        self._validate_revision_trace(run_dir, 2)
        trace_refs.append("drafts/revision_trace_v2.json")
        trace_hashes.append(self.sha256(run_dir / trace_refs[-1]))
        if version == 3:
            self._validate_revision_trace(run_dir, 3)
            trace_refs.append("drafts/revision_trace_v3.json")
            trace_hashes.append(self.sha256(run_dir / trace_refs[-1]))
        return trace_refs, trace_hashes

    def _validate_revision_trace(self, run_dir: Path, target_version: int) -> None:
        current_version = target_version - 1
        trace_ref = f"drafts/revision_trace_v{target_version}.json"
        trace_path = run_dir / trace_ref
        trace = self._load_object(trace_path, Path(trace_ref).name)
        if trace.get("from_version") not in {current_version, f"v{current_version}"}:
            raise ValueError(f"revision_trace_v{target_version} from_version 不匹配")
        if trace.get("to_version") not in {target_version, f"v{target_version}"}:
            raise ValueError(f"revision_trace_v{target_version} to_version 不匹配")
        expected_draft_ref = f"drafts/draft_v{current_version}.json"
        expected_audit_ref = f"drafts/audit_report_v{current_version}.json"
        expected_gate_ref = f"drafts/quality_gate_report_v{current_version}.json"
        if trace.get("source_draft_ref") != expected_draft_ref:
            raise ValueError(f"revision_trace_v{target_version} source_draft_ref 不匹配")
        if trace.get("source_draft_hash") != self.sha256(run_dir / expected_draft_ref):
            raise ValueError(f"revision_trace_v{target_version} source_draft_hash 不匹配")
        figure_ref = self._require_string(trace, "source_figure_slots_ref", target_version)
        figure_path = self._safe_run_relative_path(run_dir, figure_ref)
        if trace.get("source_figure_slots_hash") != self.sha256(figure_path):
            raise ValueError(
                f"revision_trace_v{target_version} source_figure_slots_hash 不匹配"
            )
        if trace.get("source_audit_report_ref") != expected_audit_ref:
            raise ValueError(f"revision_trace_v{target_version} source_audit_report_ref 不匹配")
        if trace.get("source_audit_report_hash") != self.sha256(run_dir / expected_audit_ref):
            raise ValueError(f"revision_trace_v{target_version} source_audit_report_hash 不匹配")
        if trace.get("source_quality_gate_report_ref") != expected_gate_ref:
            raise ValueError(
                f"revision_trace_v{target_version} source_quality_gate_report_ref 不匹配"
            )
        if trace.get("source_quality_gate_report_hash") != self.sha256(run_dir / expected_gate_ref):
            raise ValueError(
                f"revision_trace_v{target_version} source_quality_gate_report_hash 不匹配"
            )
        audit = DraftAuditReport.model_validate(
            self._load_object(run_dir / expected_audit_ref, Path(expected_audit_ref).name)
        )
        gate = DraftQualityGateReport.model_validate(
            self._load_object(run_dir / expected_gate_ref, Path(expected_gate_ref).name)
        )
        if figure_ref != gate.source_figure_slots_ref or figure_ref != audit.source_figure_slots_ref:
            raise ValueError(f"revision_trace_v{target_version} figure_slots lineage 不匹配")

    @staticmethod
    def _require_string(trace: dict[str, Any], key: str, target_version: int) -> str:
        value = trace.get(key)
        if not isinstance(value, str) or not value.strip():
            raise ValueError(f"revision_trace_v{target_version} {key} 缺失")
        return value

    def _validate_draft_structure_against_section_plan(
        self,
        state: DocForgeState,
        draft: dict[str, Any],
    ) -> None:
        if state.frozen_doc_plan is None or state.outline is None or not state.section_plan:
            raise ValueError("导出要求 FrozenDocPlan、Outline 和 SectionPlan")
        chapters = draft.get("chapters")
        if not isinstance(chapters, list):
            raise ValueError("draft.chapters 必须是 list")
        outline_chapters = [str(chapter.get("title", "")) for chapter in state.outline.chapters]
        draft_chapters = [str(chapter.get("title", "")) for chapter in chapters]
        locked = state.frozen_doc_plan.chapter_policy.get("locked_top_level_chapters")
        locked_chapters = [str(item) for item in locked] if isinstance(locked, list) else []
        if draft_chapters != outline_chapters:
            raise ValueError("draft chapter 顺序或标题与 Outline 不一致")
        if locked_chapters and draft_chapters != locked_chapters:
            raise ValueError("draft 一级目录与 FrozenDocPlan 不一致")

        expected_by_id = {item.section_id: item for item in state.section_plan}
        actual_sections = self._draft_sections_in_order(draft)
        actual_ids = [str(section.get("section_id", "")) for section in actual_sections]
        expected_ids = [item.section_id for item in state.section_plan]
        if actual_ids != expected_ids:
            raise ValueError("draft section 顺序或数量与 SectionPlan 不一致")
        if len(set(actual_ids)) != len(actual_ids):
            raise ValueError("draft section_id 重复")
        outline_ids = [str(node.section.get("section_id", "")) for node in iter_outline_sections(state.outline)]
        if actual_ids != outline_ids:
            raise ValueError("draft section 顺序与 Outline 不一致")
        for section in actual_sections:
            section_id = str(section.get("section_id", ""))
            plan = expected_by_id.get(section_id)
            if plan is None:
                raise ValueError("draft 包含 SectionPlan 外章节")
            if section.get("section_title") != plan.section_title:
                raise ValueError("draft section_title 与 SectionPlan 不一致")
            if section.get("section_path") != plan.section_path:
                raise ValueError("draft section_path 与 SectionPlan 不一致")

    @staticmethod
    def _draft_sections_in_order(draft: dict[str, Any]) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        for chapter in draft.get("chapters", []):
            if not isinstance(chapter, dict):
                raise ValueError("draft chapter 必须是 object")
            raw_sections = chapter.get("sections")
            if not isinstance(raw_sections, list):
                raise ValueError("draft chapter.sections 必须是 list")
            for section in raw_sections:
                if not isinstance(section, dict):
                    raise ValueError("draft section 必须是 object")
                if not str(section.get("section_id", "")).strip():
                    raise ValueError("draft section_id 缺失")
                sections.append(section)
        return sections

    def _collect_user_visible_docx_texts(
        self,
        context: ExportContext,
        state: DocForgeState,
        generated_at: str,
    ) -> list[str]:
        texts: list[str] = []
        self._append_docx_text(
            texts,
            "风险版文档" if context.export_type == "risk" else "软件著作权文档",
        )
        if context.export_type == "risk":
            self._append_docx_text(texts, RISK_NOTICE)
            self._append_docx_text(texts, f"blocker_count: {context.audit_report.summary.blocker_count}")
            self._append_docx_text(texts, f"major_count: {context.audit_report.summary.major_count}")
            self._append_docx_text(texts, "不建议直接作为正式提交版本。")
        else:
            self._append_docx_text(texts, "软件基本信息")
        for text in self._software_identity_texts(state, generated_at):
            self._append_docx_text(texts, text)

        slots_by_section = self._figure_slots_by_section(context.figures)
        for chapter in context.draft["chapters"]:
            self._append_docx_text(texts, str(chapter.get("title", "")))
            for section in chapter.get("sections", []):
                self._append_docx_text(texts, str(section.get("section_title", "")))
                self._append_docx_text(texts, str(section.get("content", "")))
                for index, slot in enumerate(
                    slots_by_section.get(section["section_id"], []),
                    start=1,
                ):
                    caption = self._figure_slot_caption(slot)
                    self._append_docx_text(texts, caption)
                    self._append_docx_text(texts, slot.get("recommended_screenshot"))
                    self._append_docx_text(texts, f"[此处建议插入：{caption}]")
                    self._append_docx_text(texts, f"图 {index} {caption}")
        return texts

    @staticmethod
    def _append_docx_text(texts: list[str], value: Any) -> None:
        if isinstance(value, str) and value:
            texts.append(value)

    def _validate_user_visible_docx_texts_safe(
        self,
        texts: list[str],
        state: DocForgeState,
        draft: dict[str, Any] | None = None,
    ) -> None:
        text = "\n".join(texts)
        lower_text = text.lower()
        if "ev_" in lower_text or EV_ID_PATTERN.search(text):
            raise ValueError("DOCX 可见文本包含内部字段，拒绝导出")
        for token in INTERNAL_FIELD_TOKENS:
            if token in lower_text:
                raise ValueError("DOCX 可见文本包含内部字段，拒绝导出")
        if self._contains_forbidden_quote_marker(text):
            raise ValueError("DOCX 可见文本包含 raw quote 字段，拒绝导出")
        for quote in self._state_raw_quotes(state):
            if quote and quote in text:
                raise ValueError("DOCX 可见文本包含 raw quote，拒绝导出")
        if draft is not None:
            self._validate_draft_citation_quotes_not_rendered(draft)

    @staticmethod
    def _contains_forbidden_quote_marker(text: str) -> bool:
        lower_text = text.lower()
        if '"quote"' in lower_text or "'quote'" in lower_text or "quote:" in lower_text:
            return True
        if "quote" not in lower_text:
            return False
        internal_context = ("evidence_id", "source_id", "finding_id", "audit_report", "raw quote")
        return any(token in lower_text for token in internal_context)

    def _validate_draft_citation_quotes_not_rendered(self, draft: dict[str, Any]) -> None:
        for section in self._draft_sections_in_order(draft):
            content = str(section.get("content", ""))
            citations = section.get("citations", [])
            if not isinstance(citations, list):
                continue
            for citation in citations:
                if not isinstance(citation, dict):
                    continue
                quote = citation.get("quote")
                if isinstance(quote, str) and quote.strip() and quote in content:
                    raise ValueError("DOCX 可见文本包含 raw quote，拒绝导出")

    @staticmethod
    def _state_raw_quotes(state: DocForgeState) -> set[str]:
        quotes: set[str] = set()
        for capability in state.product_capabilities:
            for support in capability.evidence_supports:
                if support.quote.strip():
                    quotes.add(support.quote.strip())
        for fact in state.product_facts:
            for quote in fact.supporting_quotes:
                if isinstance(quote, str) and quote.strip():
                    quotes.add(quote.strip())
        return quotes

    def _render_docx(
        self,
        context: ExportContext,
        state: DocForgeState,
        output_path: Path,
        generated_at: str,
    ) -> None:
        doc = Document()
        doc.add_heading(
            "风险版文档" if context.export_type == "risk" else "软件著作权文档",
            level=0,
        )
        if context.export_type == "risk":
            doc.add_paragraph(RISK_NOTICE)
            doc.add_paragraph(f"blocker_count: {context.audit_report.summary.blocker_count}")
            doc.add_paragraph(f"major_count: {context.audit_report.summary.major_count}")
            doc.add_paragraph("不建议直接作为正式提交版本。")
        else:
            doc.add_heading("软件基本信息", level=1)
        self._render_software_identity(doc, state, generated_at)
        slots_by_section = self._figure_slots_by_section(context.figures)
        for chapter in context.draft["chapters"]:
            doc.add_heading(str(chapter.get("title", "")), level=1)
            for section in chapter.get("sections", []):
                level = int(section.get("section_level", 2))
                doc.add_heading(str(section.get("section_title", "")), level=2 if level == 2 else 3)
                doc.add_paragraph(str(section.get("content", "")))
                for index, slot in enumerate(slots_by_section.get(section["section_id"], []), start=1):
                    caption = self._figure_slot_caption(slot)
                    doc.add_paragraph(f"[此处建议插入：{caption}]")
                    doc.add_paragraph(f"图 {index} {caption}")
        doc.save(output_path)

    def _render_software_identity(
        self,
        doc: Document,
        state: DocForgeState,
        generated_at: str,
    ) -> None:
        for text in self._software_identity_texts(state, generated_at):
            doc.add_paragraph(text)

    @staticmethod
    def _software_identity_texts(state: DocForgeState, generated_at: str) -> list[str]:
        identity = state.frozen_doc_plan.software_identity if state.frozen_doc_plan else {}
        name = identity.get("software_name") or identity.get("target_product_name") or "未命名软件"
        version = identity.get("version") or "未指定"
        target_name = identity.get("target_product_name") or name
        return [
            f"软件名称：{name}",
            f"软件版本：{version}",
            f"目标产品名称：{target_name}",
            f"文档生成时间：{generated_at}",
        ]

    @staticmethod
    def _figure_slots_by_section(figures: FigureSlotResult) -> dict[str, list[dict[str, Any]]]:
        result: dict[str, list[dict[str, Any]]] = {}
        for slot in figures.figure_slots:
            result.setdefault(slot.section_id, []).append(slot.model_dump(mode="json"))
        return result

    @staticmethod
    def _figure_slot_caption(slot: dict[str, Any]) -> str:
        caption = slot.get("recommended_caption") or slot.get("recommended_screenshot")
        if not isinstance(caption, str) or not caption.strip():
            return "待补充截图"
        return caption

    @staticmethod
    def _read_docx_visible_texts(path: Path) -> list[str]:
        doc = Document(path)
        texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        texts.append(cell.text)
        return texts

    def _build_export_manifest(
        self,
        context: ExportContext,
        docx_path: Path,
        manifest_path: Path,
    ) -> ExportManifest:
        return ExportManifest(
            export_type=context.export_type,
            draft_version=f"v{context.version}",  # type: ignore[arg-type]
            output_docx_path=self._relative_to_run_exports(docx_path),
            output_docx_hash=self.sha256(docx_path),
            source_draft_ref=context.gate_report.source_draft_ref,
            source_draft_hash=context.gate_report.source_draft_hash,
            source_figure_slots_ref=context.gate_report.source_figure_slots_ref,
            source_figure_slots_hash=context.gate_report.source_figure_slots_hash,
            source_audit_report_ref=context.gate_report.source_audit_report_path,
            source_audit_report_hash=context.gate_report.source_audit_report_hash,
            source_quality_gate_report_ref=context.gate_ref,
            source_quality_gate_report_hash=self.sha256(context.gate_path),
            source_revision_trace_refs=context.revision_trace_refs,
            source_revision_trace_hashes=context.revision_trace_hashes,
            quality_gate_passed=context.gate_report.passed,
            audit_overall_passed=context.audit_report.overall_passed,
            unresolved_blocker_count=context.audit_report.summary.blocker_count,
            unresolved_major_count=context.audit_report.summary.major_count,
            risk_notice=RISK_NOTICE if context.export_type == "risk" else None,
        )

    @staticmethod
    def _relative_to_run_exports(path: Path) -> str:
        return f"exports/{path.name}"

    def _apply_success_state(
        self,
        state: DocForgeState,
        context: ExportContext,
        docx_path: Path,
    ) -> None:
        export_result = ExportResult(
            export_type=ExportType.RISK if context.export_type == "risk" else ExportType.FINAL,
            docx_path=f"exports/{docx_path.name}",
            exported_at=datetime.now(UTC).isoformat(),
            export_notes=[
                "risk_version" if context.export_type == "risk" else "normal_version",
                f"draft_v{context.version}",
            ],
        )
        state.export_result = export_result
        state.final_doc_path = export_result.docx_path
        state.workflow_status = (
            WorkflowStatus.RISK_EXPORTED
            if context.export_type == "risk"
            else WorkflowStatus.FINAL_EXPORTED
        )
        state.next_action = NextAction.STOP

    def _safe_run_relative_path(self, run_dir: Path, relative_ref: str) -> Path:
        if not isinstance(relative_ref, str) or not relative_ref.strip():
            raise ValueError("artifact ref 缺失")
        candidate = (run_dir / relative_ref).resolve()
        resolved_run_dir = run_dir.resolve()
        if not candidate.is_relative_to(resolved_run_dir):
            raise ValueError("artifact ref 越界")
        if not candidate.exists():
            raise ValueError(f"artifact 文件不存在: {relative_ref}")
        return candidate

    @staticmethod
    def _load_object(path: Path, name: str) -> dict[str, Any]:
        if not path.exists():
            raise ValueError(f"{name} 缺失")
        try:
            value = json.loads(path.read_text(encoding="utf-8"))
        except Exception as exc:
            raise ValueError(f"{name} 无法解析") from exc
        if not isinstance(value, dict):
            raise ValueError(f"{name} 必须是 JSON 对象")
        return value

    @staticmethod
    def sha256(path: Path) -> str:
        return hashlib.sha256(path.read_bytes()).hexdigest()

    @staticmethod
    def _restore_state_file(state_path: Path, original: bytes) -> None:
        rollback = state_path.with_suffix(".json.export_rollback.tmp")
        try:
            rollback.write_bytes(original)
            rollback.replace(state_path)
        finally:
            if rollback.exists():
                rollback.unlink()

    @staticmethod
    def _snapshot_file(path: Path) -> bytes | None:
        return path.read_bytes() if path.exists() else None

    @staticmethod
    def _restore_output_file(path: Path, original: bytes | None) -> None:
        if original is None:
            if path.exists():
                path.unlink()
            return
        restore_tmp = path.with_suffix(path.suffix + ".restore.tmp")
        try:
            restore_tmp.write_bytes(original)
            restore_tmp.replace(path)
        finally:
            if restore_tmp.exists():
                restore_tmp.unlink()


def docx_has_embedded_media(path: Path) -> bool:
    """Return whether a DOCX package contains image/media parts."""
    with zipfile.ZipFile(path) as archive:
        return any(name.startswith("word/media/") for name in archive.namelist())
