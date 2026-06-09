"""DOCX parser."""

from pathlib import Path

from docx import Document

from .base import BaseParser, ParsedChunk


class DocxParser(BaseParser):
    """Parse DOCX paragraphs and tables into chunks."""

    @property
    def supported_suffixes(self) -> frozenset[str]:
        return frozenset({".docx"})

    def parse(self, path: Path) -> list[ParsedChunk]:
        self.validate_path(path)
        document = Document(str(path))
        chunks: list[ParsedChunk] = []

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            chunks.append(
                ParsedChunk(
                    text=text,
                    source=str(path),
                    metadata={"parser": "docx", "block_type": "paragraph", "index": index},
                )
            )

        for index, table in enumerate(document.tables):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            text = "\n".join(rows).strip()
            if not text:
                continue
            chunks.append(
                ParsedChunk(
                    text=text,
                    source=str(path),
                    metadata={"parser": "docx", "block_type": "table", "index": index},
                )
            )

        return chunks
